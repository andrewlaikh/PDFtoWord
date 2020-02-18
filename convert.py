from pdf2image import convert_from_path
import sys
from docx import Document
from docx.shared import Cm
import os

filename = sys.argv[1]
# Converts each page into jpg image
pages = convert_from_path(filename, size=(450, None))
print('converting ' + str(filename))
count = 1
pageList = []
for page in pages:
    pageName = 'out' + str(count) + '.jpg'
    pageList.append(pageName)
    page.save(pageName, 'JPEG')
    # print('page ' + str(count) + ' done.')
    count += 1

# Create document to paste files in.
print('creating document')
document = Document()
for pageName in pageList:
    document.add_picture(pageName, width=Cm(11.12))
    document.add_paragraph();
document.save('output.docx')

# Remove JPG files
print('deleting jpg files')
currentDirectory = os.getcwd()
dirList = os.listdir(currentDirectory)
for item in dirList:
    if item.endswith(".jpg"):
         os.remove(os.path.join(currentDirectory, item))
         
print("Done!")
